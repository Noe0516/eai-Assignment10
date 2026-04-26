"""Run all warehouse RL experiments and generate report artifacts."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Callable

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

from continuous_warehouse import ContinuousWarehouse
from discretized_q import DiscretizedConfig, train_discretized_q
from dqn_agent import DQNConfig, train_dqn
from tile_coded_q import TileQConfig, train_tile_coded_q


OUT_DIR = Path(__file__).resolve().parent / "artifacts"
OUT_DIR.mkdir(exist_ok=True)


def rolling_average(values: list[float], window: int = 100) -> np.ndarray:
    arr = np.asarray(values, dtype=np.float32)
    if arr.size < window:
        return arr
    kernel = np.ones(window, dtype=np.float32) / window
    return np.convolve(arr, kernel, mode="valid")


def vector_for_action(action: int) -> tuple[float, float]:
    if action == 0:
        return (0.0, 1.0)
    if action == 1:
        return (0.0, -1.0)
    if action == 2:
        return (1.0, 0.0)
    return (-1.0, 0.0)


def evaluate_policy(env: ContinuousWarehouse, choose_action: Callable[[np.ndarray], int], n_episodes: int = 200) -> float:
    returns = []
    for _ in range(n_episodes):
        state = env.reset()
        done = False
        total = 0.0
        while not done:
            action = choose_action(state)
            state, reward, done, _ = env.step(action)
            total += reward
        returns.append(total)
    return float(np.mean(returns))


def oracle_action(state: np.ndarray, use_6d: bool = False) -> int:
    x, y = float(state[0]), float(state[1])
    hazard = np.array([2.0, 2.0], dtype=np.float32)
    pos = np.array([x, y], dtype=np.float32)

    # Waypoint policy routes around center hazard before heading to goal.
    waypoint = np.array([0.8, 3.2], dtype=np.float32)
    goal = np.array([3.5, 3.5], dtype=np.float32)

    target = waypoint if np.linalg.norm(pos - hazard) < 1.2 and x < 2.5 else goal
    dx, dy = target - pos

    if abs(dx) > abs(dy):
        return 2 if dx >= 0 else 3
    return 0 if dy >= 0 else 1


def episodes_to_threshold(rewards: list[float], threshold: float, window: int = 100) -> int | None:
    if len(rewards) < window:
        return None
    ra = rolling_average(rewards, window=window)
    idx = np.where(ra >= threshold)[0]
    if idx.size == 0:
        return None
    return int(idx[0] + window)


def plot_learning_curves(
    curves: dict[str, list[float]],
    title: str,
    out_file: Path,
    show_plots: bool = True,
) -> None:
    plt.figure(figsize=(9, 5))
    for name, rewards in curves.items():
        plt.plot(rolling_average(rewards, 100), label=name, linewidth=2)
    plt.title(title)
    plt.xlabel("Episode")
    plt.ylabel("Rolling Avg Reward (window=100)")
    plt.grid(alpha=0.3)
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_file, dpi=150)
    if show_plots:
        plt.show()
    plt.close()


def plot_policy_field(
    env: ContinuousWarehouse,
    name: str,
    policy_fn: Callable[[np.ndarray], int],
    out_file: Path,
    show_plots: bool = True,
) -> None:
    xs = np.linspace(0.2, 3.8, 15)
    ys = np.linspace(0.2, 3.8, 15)
    X, Y = np.meshgrid(xs, ys)
    U = np.zeros_like(X)
    V = np.zeros_like(Y)

    for i in range(X.shape[0]):
        for j in range(X.shape[1]):
            state = np.array([X[i, j], Y[i, j], 0.0, 0.5], dtype=np.float32)
            if env.use_6d:
                state = np.array([X[i, j], Y[i, j], 0.0, 0.5, 0.5, 1.0], dtype=np.float32)
            action = policy_fn(state)
            u, v = vector_for_action(action)
            U[i, j] = u
            V[i, j] = v

    plt.figure(figsize=(6, 6))
    plt.quiver(X, Y, U, V, color="tab:blue", alpha=0.8)
    circle_goal = plt.Circle((3.5, 3.5), 0.5, color="green", alpha=0.2)
    circle_hazard = plt.Circle((2.0, 2.0), 0.5, color="red", alpha=0.25)
    ax = plt.gca()
    ax.add_patch(circle_goal)
    ax.add_patch(circle_hazard)
    plt.xlim(0, 4)
    plt.ylim(0, 4)
    plt.title(f"Policy Vector Field: {name}")
    plt.xlabel("x")
    plt.ylabel("y")
    plt.grid(alpha=0.2)
    plt.tight_layout()
    plt.savefig(out_file, dpi=150)
    if show_plots:
        plt.show()
    plt.close()


def generate_docx_report(metrics: dict, out_docx: Path) -> None:
    doc = Document()
    doc.add_heading("Exercise 2 Report: Continuous Warehouse RL", level=1)

    doc.add_heading("Summary Metrics (4D state)", level=2)
    doc.add_paragraph(f"Optimal reference return (oracle policy): {metrics['optimal_return']:.3f}")

    for row in metrics["comparison_rows"]:
        doc.add_paragraph(
            f"{row['method']}: episodes_to_80={row['episodes_to_80']}, "
            f"final_avg_last_200={row['final_avg_last_200']:.3f}, params={row['parameters']}"
        )

    doc.add_heading("Ablation finding", level=2)
    doc.add_paragraph(metrics["ablation_takeaway"])

    doc.add_heading("Bonus 6D finding", level=2)
    doc.add_paragraph(metrics["bonus_takeaway"])

    doc.add_heading("Figures", level=2)
    for image in metrics["images"]:
        doc.add_paragraph(Path(image).name)
        doc.add_picture(image, width=Inches(6.5))

    doc.save(out_docx)


def main(show_plots: bool = True) -> None:
    seed = 42

    # 4D baseline training.
    env_disc = ContinuousWarehouse(use_6d=False, seed=seed)
    env_tile = ContinuousWarehouse(use_6d=False, seed=seed)
    env_dqn = ContinuousWarehouse(use_6d=False, seed=seed)

    disc_cfg = DiscretizedConfig(episodes=50)
    tile_cfg = TileQConfig(episodes=50)
    dqn_cfg = DQNConfig(episodes=50)

    disc_agent, disc_rewards = train_discretized_q(
        env_disc,
        bins_per_dim=[20, 20],
        selected_dims=[0, 1],
        cfg=disc_cfg,
        seed=seed,
    )
    tile_agent, tile_rewards = train_tile_coded_q(env_tile, cfg=tile_cfg, seed=seed)
    dqn_agent, dqn_rewards = train_dqn(env_dqn, cfg=dqn_cfg, use_replay=True, use_target=True, seed=seed)

    compare_plot = OUT_DIR / "learning_curves_comparison.png"
    plot_learning_curves(
        {
            "Discretized Q (2D bins)": disc_rewards,
            "Tile-Coded Linear Q": tile_rewards,
            "DQN (full 4D state)": dqn_rewards,
        },
        "4D Warehouse: Learning Curve Comparison",
        compare_plot,
        show_plots=show_plots,
    )

    # DQN ablations.
    env_ab_full = ContinuousWarehouse(use_6d=False, seed=seed)
    env_ab_no_replay = ContinuousWarehouse(use_6d=False, seed=seed)
    env_ab_no_target = ContinuousWarehouse(use_6d=False, seed=seed)

    _, dqn_full_rewards = train_dqn(env_ab_full, cfg=dqn_cfg, use_replay=True, use_target=True, seed=seed)
    _, dqn_no_replay_rewards = train_dqn(env_ab_no_replay, cfg=dqn_cfg, use_replay=False, use_target=True, seed=seed)
    _, dqn_no_target_rewards = train_dqn(env_ab_no_target, cfg=dqn_cfg, use_replay=True, use_target=False, seed=seed)

    ablation_plot = OUT_DIR / "dqn_ablation_curves.png"
    plot_learning_curves(
        {
            "Full DQN": dqn_full_rewards,
            "No Replay": dqn_no_replay_rewards,
            "No Target Network": dqn_no_target_rewards,
        },
        "DQN Ablation Study",
        ablation_plot,
        show_plots=show_plots,
    )

    # Bonus 6D comparison.
    env_disc_6d = ContinuousWarehouse(use_6d=True, seed=seed)
    env_dqn_6d = ContinuousWarehouse(use_6d=True, seed=seed)

    disc6_agent, disc6_rewards = train_discretized_q(
        env_disc_6d,
        bins_per_dim=[10, 10, 10, 10, 10, 10],
        selected_dims=[0, 1, 2, 3, 4, 5],
        cfg=DiscretizedConfig(episodes=50, epsilon_decay=0.999, epsilon_min=0.1),
        seed=seed,
    )
    dqn6_agent, dqn6_rewards = train_dqn(
        env_dqn_6d,
        cfg=DQNConfig(episodes=50, epsilon_decay=0.998),
        use_replay=True,
        use_target=True,
        seed=seed,
    )

    bonus_plot = OUT_DIR / "bonus_6d_comparison.png"
    plot_learning_curves(
        {
            "Discretized Q (6D, 10 bins)": disc6_rewards,
            "DQN (6D)": dqn6_rewards,
        },
        "Bonus 6D: Curse of Dimensionality",
        bonus_plot,
        show_plots=show_plots,
    )

    # Policy vector fields.
    plot_policy_field(
        ContinuousWarehouse(use_6d=False, seed=seed),
        "Discretized Q",
        lambda s: int(np.argmax(disc_agent.q_values(s))),
        OUT_DIR / "policy_discretized.png",
        show_plots=show_plots,
    )
    plot_policy_field(
        ContinuousWarehouse(use_6d=False, seed=seed),
        "Tile-Coded Q",
        lambda s: int(np.argmax(tile_agent.q_values(s))),
        OUT_DIR / "policy_tile.png",
        show_plots=show_plots,
    )
    plot_policy_field(
        ContinuousWarehouse(use_6d=False, seed=seed),
        "DQN",
        lambda s: dqn_agent.select_action(s, epsilon=0.0),
        OUT_DIR / "policy_dqn.png",
        show_plots=show_plots,
    )

    # Metrics for report table.
    env_eval = ContinuousWarehouse(use_6d=False, seed=seed)
    optimal_return = evaluate_policy(env_eval, lambda s: oracle_action(s), n_episodes=200)
    threshold = 0.8 * optimal_return

    comparison_rows = [
        {
            "method": "Discretized Q-learning",
            "episodes_to_80": episodes_to_threshold(disc_rewards, threshold),
            "final_avg_last_200": float(np.mean(disc_rewards[-200:])),
            "parameters": disc_agent.n_parameters,
        },
        {
            "method": "Tile-coded linear Q-learning",
            "episodes_to_80": episodes_to_threshold(tile_rewards, threshold),
            "final_avg_last_200": float(np.mean(tile_rewards[-200:])),
            "parameters": tile_agent.n_parameters,
        },
        {
            "method": "DQN",
            "episodes_to_80": episodes_to_threshold(dqn_rewards, threshold),
            "final_avg_last_200": float(np.mean(dqn_rewards[-200:])),
            "parameters": dqn_agent.n_parameters,
        },
    ]

    visited_bins = int(np.count_nonzero(disc6_agent.visit_counts))
    total_bins = int(np.prod(disc6_agent.visit_counts.shape))
    visited_pct = 100.0 * visited_bins / total_bins

    full_tail = float(np.mean(dqn_full_rewards[-200:]))
    no_replay_tail = float(np.mean(dqn_no_replay_rewards[-200:]))
    no_target_tail = float(np.mean(dqn_no_target_rewards[-200:]))

    if abs(full_tail - no_replay_tail) > abs(full_tail - no_target_tail):
        ablation_takeaway = (
            "Removing experience replay harmed stability more than removing the target network, "
            "which matches the deadly triad intuition (function approximation + bootstrapping + off-policy updates)."
        )
    else:
        ablation_takeaway = (
            "Removing the target network had the larger impact in this run, but both components improved stability "
            "by reducing target drift and temporal correlation."
        )

    bonus_takeaway = (
        f"6D discretized table size is {10**6 * 4:,} Q-values. Only {visited_pct:.2f}% of bins were visited in training, "
        "while DQN generalized over the full 6D state with far fewer parameters."
    )

    metrics = {
        "optimal_return": optimal_return,
        "comparison_rows": comparison_rows,
        "ablation_takeaway": ablation_takeaway,
        "bonus_takeaway": bonus_takeaway,
        "images": [
            str(compare_plot),
            str(ablation_plot),
            str(bonus_plot),
            str(OUT_DIR / "policy_discretized.png"),
            str(OUT_DIR / "policy_tile.png"),
            str(OUT_DIR / "policy_dqn.png"),
        ],
    }

    report_docx = Path(__file__).resolve().parent / "exercise_2_report.docx"
    generate_docx_report(metrics, report_docx)

    # Save machine-readable metrics for quick inspection.
    metrics_txt = Path(__file__).resolve().parent / "artifacts" / "metrics_summary.txt"
    lines = [
        f"Optimal reference return: {optimal_return:.4f}",
        f"Threshold (80%): {threshold:.4f}",
        "",
        "Method, EpisodesTo80, FinalAvgLast200, Parameters",
    ]
    for row in comparison_rows:
        lines.append(
            f"{row['method']}, {row['episodes_to_80']}, {row['final_avg_last_200']:.4f}, {row['parameters']}"
        )
    lines.extend(
        [
            "",
            f"DQN tail rewards: full={full_tail:.4f}, no_replay={no_replay_tail:.4f}, no_target={no_target_tail:.4f}",
            f"6D visited bins: {visited_bins}/{total_bins} ({visited_pct:.2f}%)",
            bonus_takeaway,
            ablation_takeaway,
        ]
    )
    metrics_txt.write_text("\n".join(lines), encoding="utf-8")

    print("Artifacts written to:")
    print(OUT_DIR)
    print(report_docx)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Run warehouse RL experiments.")
    parser.add_argument(
        "--no-show",
        action="store_true",
        help="Do not open matplotlib windows; only save artifacts.",
    )
    args = parser.parse_args()
    main(show_plots=not args.no_show)
